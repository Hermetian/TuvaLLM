#!/usr/bin/env python3
"""
Align Official Transcripts (DOCX) to Rough Timestamps (MMS Output).
Generates training segments.
"""

import json
import docx
import re
from pathlib import Path
from difflib import SequenceMatcher
import unicodedata

# Paths
PROJECT_ROOT = Path(__file__).parent.parent.resolve()
DATA_DIR = PROJECT_ROOT / "data"
RAW_DIR = DATA_DIR / "raw"
PROCESSED_DIR = DATA_DIR / "processed"
TRANSCRIPTS_DIR = PROCESSED_DIR / "mms_transcripts"
SEGMENTS_DIR = PROCESSED_DIR / "segments"

SEGMENTS_DIR.mkdir(parents=True, exist_ok=True)

# Vocab for normalization check
with open(PROCESSED_DIR / "vocab.json") as f:
    VOCAB = json.load(f)

def normalize_text(text):
    """Normalize text for alignment purposes."""
    text = unicodedata.normalize("NFC", text)
    text = text.lower()
    # Remove punctuation for matching
    text = re.sub(r"[^\w\s]", "", text)
    return text.split()

def split_transcript_by_day(docx_path):
    """
    Split transcript into sections by 'Day'.
    Heuristic: Look for lines satisfying date patterns or 'Day X'.
    """
    doc = docx.Document(docx_path)
    days = {}
    current_day = "Unknown"
    
    # Heuristic patterns for day start
    # "2024 June Day 1", "Day 2", "9/12/24", etc.
    day_pattern = re.compile(r"(day\s*\d+|^\d{1,2}/\d{1,2}/\d{2,4})", re.IGNORECASE)
    
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        # Check for header
        match = day_pattern.search(text)
        if match and len(text) < 50: # Assume headers are short
            print(f"Found day header: {text}")
            # Start new day section? 
            # Need to be careful not to split on random text.
            # Let's clean the header to be a key.
            new_day = text
            if current_day != "Unknown":
                days[current_day] = " ".join(full_text)
            
            current_day = new_day
            full_text = []
        else:
            full_text.append(text)
            
    # Add last day
    if current_day != "Unknown":
        days[current_day] = " ".join(full_text)
    
    # Fallback: if no headers found, return entire text as one chunk
    if not days and full_text:
        days["all"] = " ".join(full_text)
        
    return days

def get_mms_files_for_session(session_name):
    """Find JSON transcripts matching a session."""
    # Heuristic mapping
    all_files = sorted(list(TRANSCRIPTS_DIR.glob("*.json")))
    matching = []
    
    for f in all_files:
        name = f.name.lower()
        if session_name == "june":
            if "june" in name or "24-06-24" in name:
                matching.append(f)
        elif session_name == "december":
            if "december" in name or "12-24" in name:
                matching.append(f)
                
    return matching

def align_texts(official_text, mms_words):
    """Align official text to MMS words sequence."""
    official_words = normalize_text(official_text)
    mms_word_strs = [w["word"] for w in mms_words]
    
    print(f"Aligning {len(official_words)} official words to {len(mms_words)} audio words...")
    
    matcher = SequenceMatcher(None, official_words, mms_word_strs)
    
    segments = []
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            # Match!
            # Filtering:
            # - Length > 5 words (avoid noise)
            # - Duration < 30s (if it's huge, split it - TODO)
            
            match_len = i2 - i1
            if match_len < 3: continue
            
            segment_text_words = official_text.split()[i1:i2] # Roughly map back to original text? 
            # Note: official_text.split() indices might not match normalized indices if punctuation removal shifted things.
            # To be precise, we should align normalized to normalized, but keep a mapping to original.
            # For simplicity now, let's use the normalized words (joined) or better:
            # Just rebuild from normalized + basic punctuation?
            # Or just save the normalized text. ASR training usually ignores punctuation anyway.
            
            segment_text = " ".join(official_words[i1:i2])
            
            start = mms_words[j1]["start"]
            end = mms_words[j2-1]["end"]
            duration = end - start
            
            if duration < 0.5: continue
            
            # Split long segments (>20s) heuristically at silence gaps?
            # We don't have silence info here, only word timestamps.
            # Detect gaps in mms_words? 
            # If mms_words[k+1].start - mms_words[k].end > 0.5s, split.
            
            # Sub-segmentation within the match
            current_seg_start = start
            current_seg_words = []
            current_audio_path = mms_words[j1].get("audio_path")
            
            for k in range(match_len):
                w_idx_mms = j1 + k
                w_idx_off = i1 + k
                
                word_obj = mms_words[w_idx_mms]
                word_text = official_words[w_idx_off]
                word_audio = word_obj.get("audio_path")
                
                # Check for gap, duration limit, OR file boundary
                if k > 0:
                    prev_end = mms_words[w_idx_mms-1]["end"]
                    gap = word_obj["start"] - prev_end
                    curr_dur = word_obj["end"] - current_seg_start
                    file_changed = (word_audio != current_audio_path)
                    
                    if gap > 0.5 or curr_dur > 20.0 or file_changed:
                        # Flush current segment
                        if current_seg_words and current_audio_path:
                            segments.append({
                                "text": " ".join(current_seg_words),
                                "start": current_seg_start,
                                "end": prev_end,
                                "audio_path": current_audio_path
                            })
                        # Start new
                        current_seg_start = word_obj["start"]
                        current_seg_words = []
                        current_audio_path = word_audio
                
                current_seg_words.append(word_text)
                
            # Flush final
            if current_seg_words and current_audio_path:
                segments.append({
                    "text": " ".join(current_seg_words),
                    "start": current_seg_start,
                    "end": mms_words[j2-1]["end"],
                    "audio_path": current_audio_path
                })
                
    return segments

def main():
    print("="*60)
    print("Aligning Texts...")
    
    # 1. Load Transcripts
    print("\nLoading Transcripts...")
    # June
    june_docx = next(RAW_DIR.rglob("*June*.docx"), None)
    if june_docx:
        print(f"Aligning June Session: {june_docx.name}")
        # Simplification: Treat whole file as one stream for June (assuming audio files cover it all)
        # Better: split by day.
        # Implementation Detail: For MVP, align entire June text to concatenated June Audio text.
        
        full_text = load_docx_text(june_docx)
        
        # Load all June audio transcripts
        mms_files = get_mms_files_for_session("june")
        # Load and concat words (adjusting timestamps? No, files are separate)
        # We process file-by-file alignment? No, we don't know which text belongs to which file.
        # We must concat ALL audio words (with file_id) and align to FULL text.
        
        all_audio_words = []
        for mf in mms_files:
            print(f"  Loading {mf.name}...")
            words = load_mms_transcript(mf) # Returns [{word, start, end}]
            # Add file info
            audio_fname = mf.stem + ".wav" # Assumption
            # Find matching wav in audio dir (recursive search might be needed if structure flat)
            # AUDIO_DIR is flat in preprocess_audio.py? Yes.
            audio_path = PROCESSED_DIR / "audio_16k" / audio_fname
            
            for w in words:
                w["audio_path"] = str(audio_path)
                all_audio_words.append(w)
                
        if not all_audio_words:
            print("  No MMS transcripts found for June yet.")
        else:
            segments = align_texts(full_text, all_audio_words)
            print(f"  Generated {len(segments)} segments for June.")
            
            # Save segments
            save_path = SEGMENTS_DIR / "june_segments.json"
            with open(save_path, "w") as f:
                json.dump(segments, f, indent=2)

    # December (Repeat logic)
    dec_docx = next(RAW_DIR.rglob("*Tesema*.docx"), None)
    if dec_docx:
        print(f"\nAligning December Session: {dec_docx.name}")
        full_text = load_docx_text(dec_docx)
        mms_files = get_mms_files_for_session("december")
        all_audio_words = []
        for mf in mms_files:
            print(f"  Loading {mf.name}...")
            words = load_mms_transcript(mf)
            audio_fname = mf.stem + ".wav"
            audio_path = PROCESSED_DIR / "audio_16k" / audio_fname
            for w in words:
                w["audio_path"] = str(audio_path)
                all_audio_words.append(w)
        
        if all_audio_words:
            segments = align_texts(full_text, all_audio_words)
            print(f"  Generated {len(segments)} segments for December.")
            save_path = SEGMENTS_DIR / "december_segments.json"
            with open(save_path, "w") as f:
                json.dump(segments, f, indent=2)

def load_docx_text(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return " ".join(full_text)

def load_mms_transcript(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    words = []
    for chunk in data:
        chunk_text = chunk["text"].lower()
        chunk_words = chunk_text.split()
        if not chunk_words: continue
        start = chunk["start"]
        duration = chunk["end"] - chunk["start"]
        word_dur = duration / len(chunk_words)
        for i, w in enumerate(chunk_words):
            words.append({
                "word": w,
                "start": start + i * word_dur,
                "end": start + (i+1) * word_dur
            })
    return words

if __name__ == "__main__":
    main()
