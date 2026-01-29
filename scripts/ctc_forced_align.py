#!/usr/bin/env python3
"""
CTC Forced Alignment using MMS acoustic model.
Directly aligns official transcript text to audio without ASR transcription.
"""

import json
import torch
import torchaudio
from pathlib import Path
from tqdm import tqdm
import re
import unicodedata
import docx

# Paths
PROJECT_ROOT = Path(__file__).parent.parent.resolve()
DATA_DIR = PROJECT_ROOT / "data"
RAW_DIR = DATA_DIR / "raw"
PROCESSED_DIR = DATA_DIR / "processed"
AUDIO_DIR = PROCESSED_DIR / "audio_16k"
SEGMENTS_DIR = PROCESSED_DIR / "segments_ctc"

SEGMENTS_DIR.mkdir(parents=True, exist_ok=True)

# Load vocab
with open(PROCESSED_DIR / "vocab.json") as f:
    VOCAB_DATA = json.load(f)
    VOCAB_CHARS = set(VOCAB_DATA.get("vocab", VOCAB_DATA))

# Character mapping for alignment (map unsupported chars)
CHAR_MAP = {
    'b': 'p', 'c': 'k', 'd': 't', 'j': 't', 'q': 'k',
    'w': 'u', 'x': 'k', 'y': 'i', 'z': 's',
    'ā': 'a', 'ē': 'e', 'ī': 'i', 'ū': 'u',  # Fallback if macrons cause issues
}

def normalize_for_alignment(text):
    """Normalize text for CTC alignment."""
    text = unicodedata.normalize("NFC", text)
    text = text.lower()
    # Remove punctuation except spaces
    text = re.sub(r"[^\w\s]", "", text)
    # Map unsupported characters
    result = []
    for c in text:
        if c in CHAR_MAP:
            result.append(CHAR_MAP[c])
        else:
            result.append(c)
    return "".join(result)

def load_docx_paragraphs(docx_path):
    """Load transcript as list of paragraphs with text."""
    doc = docx.Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 10:  # Skip very short lines (likely headers)
            paragraphs.append(text)
    return paragraphs

def get_audio_duration(audio_path):
    """Get audio duration in seconds."""
    import soundfile as sf
    info = sf.info(audio_path)
    return info.duration

def ctc_forced_align_chunk(model, labels, audio_chunk, transcript_chunk, device):
    """
    Perform CTC forced alignment on a chunk.
    Returns list of word segments with confidence scores.
    """
    from torchaudio.functional import forced_align
    
    # Prepare transcript
    normalized = normalize_for_alignment(transcript_chunk)
    words = transcript_chunk.lower().split()
    normalized_words = normalized.split()
    
    if not normalized.strip():
        return []
    
    # Convert to token IDs
    # labels dict maps char -> id
    tokens = []
    for c in normalized.replace(" ", "|"):  # Use | as word separator
        if c in labels:
            tokens.append(labels[c])
        # Skip unknown chars
    
    if len(tokens) < 3:
        return []
    
    tokens_tensor = torch.tensor([tokens], dtype=torch.int32, device=device)
    
    # Get emissions
    with torch.no_grad():
        emissions, _ = model(audio_chunk.unsqueeze(0).to(device))
        emissions = torch.log_softmax(emissions, dim=-1)
    
    # Forced alignment
    try:
        aligned_tokens, scores = forced_align(emissions, tokens_tensor)
    except Exception as e:
        # Alignment failed (e.g., tokens longer than frames)
        return []
    
    # Convert frame indices to time
    # MMS uses 20ms frames with stride
    frame_rate = 50  # 50 frames per second (20ms each)
    
    # Group by words (split at |)
    segments = []
    current_word_idx = 0
    current_word_start = None
    current_word_scores = []
    
    aligned_list = aligned_tokens[0].tolist()
    scores_list = scores[0].tolist()
    
    for frame_idx, (token_id, score) in enumerate(zip(aligned_list, scores_list)):
        if token_id == labels.get("|", -1):  # Word separator
            if current_word_start is not None and current_word_idx < len(words):
                # End of word
                end_time = frame_idx / frame_rate
                avg_score = sum(current_word_scores) / len(current_word_scores) if current_word_scores else 0
                segments.append({
                    "word": words[current_word_idx],
                    "start": current_word_start,
                    "end": end_time,
                    "score": avg_score
                })
                current_word_idx += 1
            current_word_start = None
            current_word_scores = []
        else:
            if current_word_start is None:
                current_word_start = frame_idx / frame_rate
            current_word_scores.append(score)
    
    # Handle last word
    if current_word_start is not None and current_word_idx < len(words):
        end_time = len(aligned_list) / frame_rate
        avg_score = sum(current_word_scores) / len(current_word_scores) if current_word_scores else 0
        segments.append({
            "word": words[current_word_idx],
            "start": current_word_start,
            "end": end_time,
            "score": avg_score
        })
    
    return segments

def process_audio_file(audio_path, transcript_text, model, labels, device):
    """
    Process a single audio file with its transcript.
    Returns high-confidence segments.
    """
    import soundfile as sf
    
    # Load audio
    audio_np, sr = sf.read(audio_path, dtype="float32")
    if len(audio_np.shape) > 1:
        audio_np = audio_np.mean(axis=1)
    audio = torch.tensor(audio_np)
    
    if sr != 16000:
        audio = torchaudio.functional.resample(audio, sr, 16000)
    
    total_duration = len(audio) / 16000
    
    # Normalize transcript
    paragraphs = transcript_text.split("\n")
    full_text = " ".join(paragraphs)
    
    # Process in chunks (60 seconds with overlap)
    CHUNK_SEC = 60.0
    STRIDE_SEC = 30.0
    chunk_samples = int(CHUNK_SEC * 16000)
    stride_samples = int(STRIDE_SEC * 16000)
    
    all_segments = []
    
    # Simple approach: align full text to full audio
    # For very long audio, we'd need smarter chunking
    
    # For now, let's just try direct alignment with a confidence threshold
    words = full_text.lower().split()
    
    # Split text into manageable chunks (~100 words)
    WORDS_PER_CHUNK = 50
    text_chunks = []
    for i in range(0, len(words), WORDS_PER_CHUNK):
        chunk_words = words[i:i+WORDS_PER_CHUNK]
        text_chunks.append(" ".join(chunk_words))
    
    # Estimate time per text chunk based on audio duration
    time_per_chunk = total_duration / max(len(text_chunks), 1)
    
    for chunk_idx, text_chunk in enumerate(text_chunks):
        # Estimate audio region for this text
        start_time = chunk_idx * time_per_chunk
        end_time = min((chunk_idx + 1) * time_per_chunk, total_duration)
        
        start_sample = int(start_time * 16000)
        end_sample = int(end_time * 16000)
        
        audio_chunk = audio[start_sample:end_sample]
        
        if len(audio_chunk) < 16000:  # Skip very short chunks
            continue
        
        try:
            word_segments = ctc_forced_align_chunk(
                model, labels, audio_chunk, text_chunk, device
            )
            
            # Adjust timestamps to absolute time
            for seg in word_segments:
                seg["start"] += start_time
                seg["end"] += start_time
                seg["audio_path"] = str(audio_path)
            
            all_segments.extend(word_segments)
            
        except Exception as e:
            print(f"  Chunk {chunk_idx} failed: {e}")
            continue
    
    return all_segments

def merge_word_segments(word_segments, min_words=3, max_duration=20.0, min_confidence=-5.0):
    """
    Merge consecutive word segments into utterance segments.
    Filter by confidence score.
    """
    if not word_segments:
        return []
    
    # Filter by confidence
    good_words = [w for w in word_segments if w["score"] > min_confidence]
    
    if len(good_words) < min_words:
        return []
    
    # Merge into utterances (split at gaps > 0.5s or low confidence)
    utterances = []
    current_utterance = []
    current_audio = good_words[0].get("audio_path")
    
    for i, word in enumerate(good_words):
        if not current_utterance:
            current_utterance.append(word)
            continue
        
        gap = word["start"] - current_utterance[-1]["end"]
        audio_changed = word.get("audio_path") != current_audio
        duration = word["end"] - current_utterance[0]["start"]
        
        if gap > 0.5 or audio_changed or duration > max_duration:
            # Flush current utterance
            if len(current_utterance) >= min_words:
                utterances.append({
                    "text": " ".join(w["word"] for w in current_utterance),
                    "start": current_utterance[0]["start"],
                    "end": current_utterance[-1]["end"],
                    "audio_path": current_audio,
                    "avg_score": sum(w["score"] for w in current_utterance) / len(current_utterance)
                })
            current_utterance = [word]
            current_audio = word.get("audio_path")
        else:
            current_utterance.append(word)
    
    # Flush final
    if len(current_utterance) >= min_words:
        utterances.append({
            "text": " ".join(w["word"] for w in current_utterance),
            "start": current_utterance[0]["start"],
            "end": current_utterance[-1]["end"],
            "audio_path": current_audio,
            "avg_score": sum(w["score"] for w in current_utterance) / len(current_utterance)
        })
    
    return utterances

def main():
    print("=" * 60)
    print("CTC Forced Alignment")
    print("=" * 60)
    
    device = "cpu"  # MPS has issues with some ops, use CPU for reliability
    print(f"Device: {device}")
    
    # Load MMS model for forced alignment
    print("Loading MMS model...")
    bundle = torchaudio.pipelines.MMS_FA
    model = bundle.get_model().to(device)
    model.eval()
    
    # Get labels (char -> id mapping)
    labels = bundle.get_labels()
    labels_dict = {c: i for i, c in enumerate(labels)}
    
    print(f"Labels: {len(labels)} characters")
    print(f"Sample: {labels[:20]}")
    
    # Add word separator if not present
    if "|" not in labels_dict:
        labels_dict["|"] = len(labels_dict)
    
    # Load transcripts
    june_docx = next(RAW_DIR.rglob("*June*.docx"), None)
    dec_docx = next(RAW_DIR.rglob("*Tesema*.docx"), None)
    
    all_segments = []
    
    # Process June session
    if june_docx:
        print(f"\nProcessing June session...")
        june_text = load_docx_text(june_docx)
        
        # Find June audio files
        june_audios = sorted([
            f for f in AUDIO_DIR.glob("*.wav")
            if "june" in f.name.lower() or "24-06-24" in f.name
        ])
        
        print(f"  Found {len(june_audios)} audio files")
        
        # For each audio, try to align a portion of the transcript
        # This is a simplification - ideally we'd match by day
        text_per_file = len(june_text) // max(len(june_audios), 1)
        
        for i, audio_file in enumerate(tqdm(june_audios, desc="June files")):
            # Rough text assignment
            start_char = i * text_per_file
            end_char = (i + 1) * text_per_file
            file_text = june_text[start_char:end_char]
            
            try:
                segments = process_audio_file(
                    audio_file, file_text, model, labels_dict, device
                )
                utterances = merge_word_segments(segments)
                all_segments.extend(utterances)
                print(f"    {audio_file.name}: {len(utterances)} utterances")
            except Exception as e:
                print(f"    {audio_file.name}: FAILED - {e}")
    
    # Process December session
    if dec_docx:
        print(f"\nProcessing December session...")
        dec_text = load_docx_text(dec_docx)
        
        dec_audios = sorted([
            f for f in AUDIO_DIR.glob("*.wav")
            if "december" in f.name.lower() or "12-24" in f.name
        ])
        
        print(f"  Found {len(dec_audios)} audio files")
        
        text_per_file = len(dec_text) // max(len(dec_audios), 1)
        
        for i, audio_file in enumerate(tqdm(dec_audios, desc="Dec files")):
            start_char = i * text_per_file
            end_char = (i + 1) * text_per_file
            file_text = dec_text[start_char:end_char]
            
            try:
                segments = process_audio_file(
                    audio_file, file_text, model, labels_dict, device
                )
                utterances = merge_word_segments(segments)
                all_segments.extend(utterances)
                print(f"    {audio_file.name}: {len(utterances)} utterances")
            except Exception as e:
                print(f"    {audio_file.name}: FAILED - {e}")
    
    # Save segments
    print(f"\nTotal segments: {len(all_segments)}")
    total_duration = sum(s["end"] - s["start"] for s in all_segments)
    print(f"Total duration: {total_duration/3600:.2f} hours")
    
    save_path = SEGMENTS_DIR / "ctc_aligned_segments.json"
    with open(save_path, "w") as f:
        json.dump(all_segments, f, indent=2, ensure_ascii=False)
    
    print(f"Saved to {save_path}")

def load_docx_text(docx_path):
    """Load full text from DOCX."""
    doc = docx.Document(docx_path)
    return " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

if __name__ == "__main__":
    main()
